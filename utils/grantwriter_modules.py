"""
GrantWriter AI — rubric + coverage + DOCX export helpers
Drop this file into your app (e.g., `utils/grantwriter_modules.py`) and import.

Requires: pip install python-docx
Optional (nice to have): pip install rapidfuzz (for better keyword matching) — handled gracefully if missing.
"""
from __future__ import annotations

import re
from dataclasses import dataclass, asdict
from typing import List, Dict, Tuple, Optional
from collections import Counter, defaultdict

# --- Optional fuzzy matching ---
try:
    from rapidfuzz import fuzz
    _HAS_FUZZ = True
except Exception:  # pragma: no cover
    _HAS_FUZZ = False

# ------------------------------
# 1) RUBRIC BUILDER
# ------------------------------

STOPWORDS = set(
    """
    a an and are as at be by for from has have i in is it its of on or our that the to we with your you this those these their there which will would should could can may might such if then than while where when who whom whose into upon about above below over under again further do does did not no nor only own same so too very more most less least each other any both few many much being been were was he she they them his her theirs ours yours mine me my our also via per vs versus include including like e.g eg etc i.e ie vs.
    """.split()
)

# Common council/funder criteria buckets -> canonical names + seed terms
CRITERIA_CANON = {
    "Need": ["need", "problem", "gap", "evidence", "data", "baseline"],
    "Outcomes": ["outcomes", "results", "impact", "kpi", "targets", "benefit"],
    "Community Benefit": ["community", "benefit", "equity", "inclusion", "priority groups", "co-design"],
    "Activities": ["activities", "delivery", "implementation", "work plan", "work packages"],
    "Evaluation": ["evaluation", "monitoring", "measure", "baseline", "methods", "survey"],
    "Budget": ["budget", "value for money", "costs", "co-funding", "in-kind"],
    "Risk": ["risk", "mitigation", "safeguard", "governance", "compliance"],
    "Timeline": ["timeline", "milestones", "schedule", "gantt", "phases"],
    "Objectives": ["objective", "smart", "goal"],
    "Partnerships": ["partnership", "partner", "governance", "roles", "stakeholder"],
}

# Map canonical criteria to likely UI sections
CRITERIA_TO_SECTIONS = {
    "Need": ["Problem / Need"],
    "Outcomes": ["Expected Outcomes (KPIs/metrics)", "Objectives (bullets OK)"],
    "Community Benefit": ["Target Audience", "Executive Summary"],
    "Activities": ["Activities & Delivery Plan"],
    "Evaluation": ["Evaluation (how you'll measure)"],
    "Budget": ["Budget (summary + justification)"],
    "Risk": ["Risks & Mitigation"],
    "Timeline": ["High-level Timeline"],
    "Objectives": ["Objectives (bullets OK)"],
    "Partnerships": ["Partners & Governance"],
}

@dataclass
class RubricItem:
    criterion: str
    weight: float  # 0..1 (normalized)
    sections: List[str]
    raw_text: str
    keywords: List[str]

    def to_dict(self):
        return asdict(self)

PCT_RE = re.compile(r"(\d{1,3})\s*%")
SCORE_RE = re.compile(r"(?:score|worth|points?)\s*[:\-]?\s*(\d+)", re.I)


def _clean_tokens(line: str) -> List[str]:
    words = [w.lower() for w in re.findall(r"[a-zA-Z][a-zA-Z\-]{2,}", line)]
    return [w for w in words if w not in STOPWORDS]


def _guess_canonical(label: str) -> str:
    label_l = label.lower()
    best = (None, 0)
    for canon, seeds in CRITERIA_CANON.items():
        score = 0
        for s in seeds + [canon.lower()]:
            if _HAS_FUZZ:
                score = max(score, fuzz.partial_ratio(label_l, s))
            else:
                if s in label_l:
                    score = max(score, 100)
        if score > best[1]:
            best = (canon, score)
    return best[0] or label.strip().title()


def build_rubric(criteria_text: str) -> List[RubricItem]:
    """Parse a funder’s criteria (verbatim) into a weighted rubric.

    Rules:
      - Detect explicit weights via percentages (e.g., "Need 30%")
      - Else detect points (e.g., "Worth 15 points")
      - Else equal weights
      - Map to likely sections
      - Extract keywords (non-stopword tokens + canonical seeds)
    """
    # Split into bullet-like lines
    lines = [l.strip(" -•\t") for l in criteria_text.splitlines() if l.strip()]
    items_tmp: List[Tuple[str, float, str]] = []  # (label, raw_weight, raw_text)

    for l in lines:
        pct = PCT_RE.search(l)
        if pct:
            items_tmp.append((l, float(pct.group(1)), l))
            continue
        pts = SCORE_RE.search(l)
        if pts:
            items_tmp.append((l, float(pts.group(1)), l))
            continue
        # fallback weight placeholder 1.0
        items_tmp.append((l, 1.0, l))

    # Normalize weights
    total = sum(w for _, w, _ in items_tmp) or 1.0
    items: List[RubricItem] = []

    for label, w, raw in items_tmp:
        # derive a short label (before dash/colon)
        head = re.split(r"[:\-]", label, maxsplit=1)[0]
        canon = _guess_canonical(head)

        kws = list({*(_clean_tokens(label)), *CRITERIA_CANON.get(canon, [])})
        sections = CRITERIA_TO_SECTIONS.get(canon, [canon])

        items.append(
            RubricItem(
                criterion=canon,
                weight=(w / total),
                sections=sections,
                raw_text=raw,
                keywords=kws,
            )
        )

    # Merge duplicates by canonical name
    merged: Dict[str, RubricItem] = {}
    for it in items:
        if it.criterion in merged:
            m = merged[it.criterion]
            m.weight += it.weight
            m.raw_text += "\n" + it.raw_text
            m.keywords = list(sorted(set(m.keywords + it.keywords)))
        else:
            merged[it.criterion] = it

    # Re-normalize after merge
    total2 = sum(m.weight for m in merged.values()) or 1.0
    for m in merged.values():
        m.weight /= total2

    # Order by weight desc
    return sorted(merged.values(), key=lambda x: x.weight, reverse=True)


# ------------------------------
# 2) COVERAGE + GAP CHECKS
# ------------------------------

NUM_RE = re.compile(r"\b\d+(?:[\.,]\d+)?\b")
MONEY_RE = re.compile(r"\$\s?\d")
PCT_WORDS = {"percent", "%"}


def coverage_score(section_text: str, rubric_item: RubricItem) -> Tuple[float, Dict[str, float]]:
    """Return coverage 0..1 plus per-keyword match scores.

    Heuristic: keyword presence (1), or fuzzy >= 70 → 0.6. A few seed keywords
    carry more weight automatically.
    """
    text = section_text.lower()
    kw_weights: Dict[str, float] = {}
    # Seed weights: canonical seeds weigh 2x
    seedset = set(CRITERIA_CANON.get(rubric_item.criterion, []))
    for kw in rubric_item.keywords:
        w = 2.0 if kw in seedset else 1.0
        score = 0.0
        if kw in text:
            score = 1.0
        elif _HAS_FUZZ:
            score = 0.6 if fuzz.partial_ratio(kw, text) >= 70 else 0.0
        kw_weights[kw] = score * w

    max_possible = sum(2.0 if k in seedset else 1.0 for k in rubric_item.keywords) or 1.0
    total_score = sum(kw_weights.values()) / max_possible
    return max(0.0, min(1.0, total_score)), kw_weights


def gap_hints(section_text: str, rubric_item: RubricItem) -> List[str]:
    t = section_text.lower()
    hints: List[str] = []

    if rubric_item.criterion in ("Outcomes", "Objectives"):
        if not NUM_RE.search(t) and not any(w in t for w in PCT_WORDS):
            hints.append("Add at least one quantified KPI (number or %).")
        if "baseline" not in t:
            hints.append("State baseline and target for each KPI.")
        if rubric_item.criterion == "Objectives" and not any(x in t for x in ["by ", "within "]):
            hints.append("Make timelines explicit (e.g., 'by Q4 2025').")

    if rubric_item.criterion == "Evaluation":
        if "method" not in t and "survey" not in t and "interview" not in t:
            hints.append("Name evaluation methods (survey/interviews/admin data).")
        if "report" not in t and "cadence" not in t and "quarter" not in t:
            hints.append("Say how often you will report (e.g., quarterly).")

    if rubric_item.criterion == "Budget":
        if not MONEY_RE.search(section_text):
            hints.append("Include dollar amounts and totals.")
        if "co-fund" not in t and "in-kind" not in t and "match" not in t:
            hints.append("Mention co-funding or in-kind support if applicable.")

    if rubric_item.criterion == "Need":
        if not NUM_RE.search(t):
            hints.append("Cite 1–2 data points demonstrating the problem size.")
        if "source" not in t and "abs" not in t and "census" not in t:
            hints.append("Reference a data source (e.g., ABS, council data).")

    if rubric_item.criterion == "Risk":
        if "likelihood" not in t or "impact" not in t:
            hints.append("Include likelihood and impact for top risks.")

    if rubric_item.criterion == "Community Benefit":
        if all(w not in t for w in ["equity", "priority", "disadvantaged", "inclusive"]):
            hints.append("Address equity/priority groups explicitly.")

    return hints


def assess_project(project_sections: Dict[str, str], rubric: List[RubricItem]) -> Dict[str, Dict]:
    """Compute weighted coverage per criterion and overall score."""
    results: Dict[str, Dict] = {}
    overall = 0.0
    for item in rubric:
        # concatenate all mapped sections
        text = "\n\n".join(project_sections.get(s, "") for s in item.sections)
        cov, kw = coverage_score(text, item)
        hints = gap_hints(text, item)
        score = cov * item.weight
        overall += score
        results[item.criterion] = {
            "weight": item.weight,
            "coverage": cov,
            "weighted": score,
            "keywords": kw,
            "sections": item.sections,
            "hints": hints,
        }
    results["__overall__"] = {"score": round(overall, 4)}
    return results


# ------------------------------
# 3) DOCX EXPORT (portal-friendly)
# ------------------------------
from docx import Document
from docx.shared import Pt, Inches
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.style import WD_STYLE_TYPE

DEFAULT_EXPORT_ORDER = [
    "Executive Summary",
    "Statement of Need",
    "Project Objectives",
    "Activities & Delivery Plan",
    "Expected Outcomes (KPIs/metrics)",
    "Evaluation (how you'll measure)",
    "Community Benefit",
    "Budget & Justification",
    "Risk Management",
    "Project Timeline",
    "Partnerships & Governance",
]

SECTION_TITLE_MAP = {
    "Problem / Need": "Statement of Need",
    "Budget (summary + justification)": "Budget & Justification",
    "High-level Timeline": "Project Timeline",
    "Risks & Mitigation": "Risk Management",
    "Objectives (bullets OK)": "Project Objectives",
    "Activities & Delivery Plan": "Activities & Delivery Plan",
    "Expected Outcomes (KPIs/metrics)": "Expected Outcomes (KPIs/metrics)",
    "Evaluation (how you'll measure)": "Evaluation (how you'll measure)",
    "Target Audience": "Community Benefit",
    "Partners & Governance": "Partnerships & Governance",
    "Executive Summary": "Executive Summary",
}


def _ensure_styles(doc: Document):
    styles = doc.styles
    # Heading 1 style
    if "Grant Heading" not in styles:
        h = styles.add_style("Grant Heading", WD_STYLE_TYPE.PARAGRAPH)
        h.base_style = styles["Heading 1"]
        h.font.size = Pt(14)
        h.font.bold = True
    if "Grant Body" not in styles:
        b = styles.add_style("Grant Body", WD_STYLE_TYPE.PARAGRAPH)
        b.base_style = styles["Normal"]
        b.font.size = Pt(11)


def export_docx(
    project: Dict[str, str],
    filename: str = "GrantWriter_Export.docx",
    order: Optional[List[str]] = None,
    meta: Optional[Dict[str, str]] = None,
    rubric_result: Optional[Dict[str, Dict]] = None,
) -> str:
    """Export the project sections to a clean DOCX.

    - `project`: dict of section name -> text
    - `order`: list of section titles in desired order (fallback to DEFAULT_EXPORT_ORDER)
    - `meta`: {title, funder, applicant, amount}
    - `rubric_result`: output of assess_project(), to include a short summary page
    Returns the path to the saved file.
    """
    doc = Document()
    _ensure_styles(doc)

    # Cover
    title = (meta or {}).get("title", project.get("Project Title", "Grant Application"))
    doc.add_paragraph(title, style="Grant Heading").alignment = WD_ALIGN_PARAGRAPH.CENTER
    subtitle_bits = []
    if meta:
        if meta.get("funder"): subtitle_bits.append(f"Funder: {meta['funder']}")
        if meta.get("applicant"): subtitle_bits.append(f"Applicant: {meta['applicant']}")
        if meta.get("amount"): subtitle_bits.append(f"Amount requested: {meta['amount']}")
    if subtitle_bits:
        doc.add_paragraph(" | ".join(subtitle_bits)).alignment = WD_ALIGN_PARAGRAPH.CENTER
    doc.add_paragraph("")

    # Optional rubric summary
    if rubric_result:
        doc.add_paragraph("Assessment Summary", style="Grant Heading")
        overall = rubric_result.get("__overall__", {}).get("score", 0)
        doc.add_paragraph(f"Overall coverage score (weighted): {round(overall*100):d}%", style="Grant Body")
        t = doc.add_table(rows=1, cols=4)
        hdr = t.rows[0].cells
        hdr[0].text = "Criterion"; hdr[1].text = "Weight"; hdr[2].text = "Coverage"; hdr[3].text = "Hints"
        for k, v in rubric_result.items():
            if k.startswith("__"): continue
            row = t.add_row().cells
            row[0].text = k
            row[1].text = f"{round(v['weight']*100):d}%"
            row[2].text = f"{round(v['coverage']*100):d}%"
            row[3].text = ("; ".join(v.get("hints", [])[:3]))
        doc.add_paragraph("")

    # Body sections
    section_order = order or DEFAULT_EXPORT_ORDER

    # Map any project keys to export-friendly titles
    normalized: Dict[str, str] = {}
    for k, v in project.items():
        mapped = SECTION_TITLE_MAP.get(k, k)
        normalized[mapped] = v

    for sec in section_order:
        content = normalized.get(sec)
        if not content:
            continue
        doc.add_paragraph(sec, style="Grant Heading")
        # Split on double newlines to keep some structure
        for para in re.split(r"\n\n+", content.strip()):
            doc.add_paragraph(para, style="Grant Body")
        doc.add_paragraph("")

    doc.save(filename)
    return filename


# ------------------------------
# 4) EXAMPLE INTEGRATION (minimal)
# ------------------------------
if __name__ == "__main__":
    sample_criteria = """
    Demonstrated Need (30%) – evidence of problem, baseline data.
    Outcomes & Impact (25%) – clear, measurable KPIs.
    Project Delivery (20%) – activities and timeline.
    Community Benefit (15%) – equity & inclusion.
    Value for Money (10%) – budget, co-funding.
    """

    rubric = build_rubric(sample_criteria)
    print("Rubric: ")
    for r in rubric:
        print(r.to_dict())

    project = {
        "Project Title": "Safer Streets: Night Lighting for 12 Estates",
        "Problem / Need": "Police callouts after dark are 28% higher than the metro average (VicPol 2024). Residents report fear of walking at night. Baseline: 42 incidents/quarter across target estates.",
        "Expected Outcomes (KPIs/metrics)": "By Q4 2025, reduce after-dark incidents by 25% (to ≤31/quarter). 300 residents attend safety workshops; 85% report improved perception of safety.",
        "Activities & Delivery Plan": "Install 180 LED fixtures; co-design placements with residents; monthly checks; safety workshops with local partners.",
        "Evaluation (how you'll measure)": "Compare incident reports vs baseline; quarterly resident survey (n≥150); asset uptime logs; publish quarterly report.",
        "Target Audience": "2,400 residents across 12 social-housing estates; priority groups include women, youth, and older residents.",
        "Budget (summary + justification)": "$180,000 LED fixtures and installation; $12,000 workshops; $8,000 evaluation; in-kind: council electricians (0.2 FTE).",
        "Risks & Mitigation": "Supply delays (Likely/Medium) → pre-procure; Vandalism (Possible/High) → tamper-proof fittings; Complaints about light spill → resident co-design.",
        "High-level Timeline": "Q1: design & permits; Q2: procurement; Q3: installation; Q4: evaluation.",
        "Partners & Governance": "Housing provider (MOU), VicPol (data), Resident groups (co-design).",
        "Executive Summary": "We will install 180 LED lights across 12 estates to reduce after-dark incidents by 25% within 12 months, benefiting 2,400 residents. Partners include Housing Victoria and VicPol.",
        "Objectives (bullets OK)": "1) Cut incidents −25% by Q4 2025; 2) 85% of participants report improved safety; 3) ≥95% asset uptime.",
    }

    result = assess_project(project, rubric)
    print("\nCoverage results:")
    for k, v in result.items():
        print(k, v)

    path = export_docx(project, filename="GrantWriter_Export_Demo.docx", meta={"funder":"City Grants 2025","applicant":"Your Council","amount":"$200,000"}, rubric_result=result)
    print(f"\nExported to: {path}")
