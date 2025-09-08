import os, io, json, datetime
from dataclasses import dataclass, asdict
from typing import Dict, List
import streamlit as st
from dotenv import load_dotenv
from docx import Document
from docx.shared import Pt, Inches
from docx.enum.text import WD_ALIGN_PARAGRAPH

# --- OpenAI client (v1) ---
try:
    from openai import OpenAI
except ImportError:
    st.write("Install deps first: pip install -r requirements.txt")
    st.stop()

# ===== ENV =====
load_dotenv()
OPENAI_API_KEY = os.getenv("OPENAI_API_KEY", "")
COUNCIL_NAME = os.getenv("COUNCIL_NAME", "Your Council")
COUNCIL_SHORT = os.getenv("COUNCIL_SHORT", "council")
BRAND_PRIMARY = os.getenv("COUNCIL_BRAND_PRIMARY", "#0D47A1")
BRAND_SECONDARY = os.getenv("COUNCIL_BRAND_SECONDARY", "#1976D2")
if not OPENAI_API_KEY:
    st.error("Missing OPENAI_API_KEY in .env")
    st.stop()
client = OpenAI(api_key=OPENAI_API_KEY)

st.set_page_config(page_title=f"GrantWriter AI – {COUNCIL_NAME}", page_icon="📝", layout="wide")
st.markdown(
    f"""
    <style>
    :root {{ --brand-primary: {BRAND_PRIMARY}; --brand-secondary: {BRAND_SECONDARY}; }}
    .brand-card {{ border:1px solid #e5e7eb; padding:14px; border-radius:14px; background:#fff; box-shadow:0 2px 16px #0000000f; }}
    .app-title {{ font-size:1.8rem; font-weight:700; margin-bottom:.25rem; }}
    .app-subtitle {{ color:#555; margin-bottom:1rem; }}
    .pill {{ border:1px solid #e5e7eb; border-radius:999px; padding:4px 10px; font-size:.8rem; margin-right:6px; }}
    .section-title {{ font-weight:700; margin:.5rem 0 .25rem; font-size:1.05rem; }}
    .small-muted {{ color:#6b7280; font-size:.85rem; }}
    </style>
    """,
    unsafe_allow_html=True,
)

# ===== DATA =====
@dataclass
class Project:
    title: str = ""
    summary: str = ""
    need: str = ""
    audience: str = ""
    objectives: str = ""
    activities: str = ""
    outcomes: str = ""
    evaluation: str = ""
    budget: str = ""
    risks: str = ""
    timeline: str = ""
    partners: str = ""

SECTIONS: List[str] = [
    "Executive Summary","Statement of Need","Project Objectives","Activities & Delivery Plan",
    "Outcomes & Evaluation","Community Benefit & Equity","Budget & Justification","Risk Management",
    "Project Timeline","Partnerships & Governance"
]

def sys_prompt() -> str:
    return f"""You are a senior grant writer for {COUNCIL_NAME} in Australia.
Write precise, persuasive, criteria-aligned responses in clear, formal English.
Use measurable, implementable actions. Include equity/access where relevant."""

def draft_prompt(section: str, p: Project, criteria: str) -> str:
    return f"""
SECTION: {section}

GRANT CRITERIA (verbatim):
{criteria.strip() or "N/A"}

PROJECT DETAILS:
- Title: {p.title}
- Summary: {p.summary}
- Problem/Need: {p.need}
- Target Audience: {p.audience}
- Objectives: {p.objectives}
- Activities: {p.activities}
- Outcomes: {p.outcomes}
- Evaluation: {p.evaluation}
- Budget: {p.budget}
- Risks & Mitigation: {p.risks}
- Timeline: {p.timeline}
- Partnerships/Governance: {p.partners}

INSTRUCTIONS:
1) Draft '{section}' in 150–250 words.
2) Align tightly to CRITERIA; mirror sub-points where present.
3) Use plain language and measurable outcomes.
4) If key info is missing, draft best-possible and append [INFO NEEDED: ...].
"""

def score_prompt(text: str, criteria: str) -> str:
    return f"""
You are an expert assessor. Score 0–100 vs criteria.

CRITERIA:
{criteria.strip() or "N/A"}

DRAFT:
{text}

Return strict JSON:
{{"score": <int>,"strengths":["..."],"gaps":["..."],"suggestions":["..."]}}
"""

def gpt_text(u: str, model: str = "gpt-4o-mini") -> str:
    r = client.chat.completions.create(
        model=model,
        messages=[{"role":"system","content":sys_prompt()},{"role":"user","content":u}],
        temperature=0.3, max_tokens=700,
    )
    return r.choices[0].message.content.strip()

def gpt_json(u: str, model: str = "gpt-4o-mini") -> Dict:
    r = client.chat.completions.create(
        model=model,
        messages=[{"role":"system","content":"Output strict JSON only."},{"role":"user","content":u}],
        temperature=0.0, response_format={"type":"json_object"}, max_tokens=700,
    )
    import json
    try:
        return json.loads(r.choices[0].message.content)
    except Exception:
        return {"score":0,"strengths":[],"gaps":["Parse error"],"suggestions":[]}

def build_docx(sections: Dict[str,str], p: Project, logo_bytes: bytes | None) -> bytes:
    doc = Document()
    tmp = None
    if logo_bytes:
        tmp = "~tmp_logo.png"
        with open(tmp,"wb") as f: f.write(logo_bytes)
        try:
            doc.add_picture(tmp, width=Inches(1.2))
            doc.paragraphs[-1].alignment = WD_ALIGN_PARAGRAPH.RIGHT
        except Exception: pass
        try: os.remove(tmp)
        except Exception: pass

    t = doc.add_paragraph(); r = t.add_run(f"{COUNCIL_NAME} – Grant Application Draft")
    r.bold=True; r.font.size = Pt(18)
    meta = doc.add_paragraph()
    meta.add_run("Project: ").bold=True; meta.add_run(p.title + "\n")
    meta.add_run("Generated: ").bold=True; meta.add_run(datetime.datetime.now().strftime("%Y-%m-%d %H:%M"))
    doc.add_paragraph("")

    for k,v in sections.items():
        h = doc.add_paragraph(); rh = h.add_run(k); rh.bold=True; rh.font.size=Pt(14)
        doc.add_paragraph((v or "").strip()); doc.add_paragraph("")
    bio = io.BytesIO(); doc.save(bio); bio.seek(0); return bio.read()

def markdown_bytes(sections: Dict[str,str]) -> bytes:
    md = "# " + COUNCIL_NAME + " – Grant Application Draft\n\n"
    for k,v in sections.items(): md += f"## {k}\n{(v or '').strip()}\n\n"
    return md.encode("utf-8")

def load_logo_bytes() -> bytes | None:
    path = os.path.join("assets","council_logo.png")
    if os.path.exists(path):
        with open(path,"rb") as f: return f.read()
    return None

# ===== HEADER =====
st.markdown(f"""
<div class="brand-card">
  <div class="app-title">📝 GrantWriter AI</div>
  <div class="app-subtitle">Draft, score, and export grant applications for <b>{COUNCIL_NAME}</b>.</div>
  <span class="pill">Drafting</span><span class="pill">Scoring</span><span class="pill">Gap checks</span><span class="pill">DOCX export</span><span class="pill">Save/Load</span>
</div>
""", unsafe_allow_html=True)

# ===== SIDEBAR =====
with st.sidebar:
    st.header("⚙️ Settings")
    use_sections = st.multiselect("Sections to generate", SECTIONS, default=SECTIONS)
    st.subheader("📋 Grant Criteria")
    criteria = st.text_area("Paste selection criteria (verbatim):", height=220)
    model = st.selectbox("Model", ["gpt-4o-mini","gpt-4o"], index=0)

    st.markdown("---")
    st.subheader("💾 Save / Load Project")
    save_name = st.text_input("Save as (filename.json)")
    if st.button("Save Project"):
        proj = st.session_state.get("project", Project())
        drafts = st.session_state.get("application", {})
        blob = {"project": asdict(proj), "criteria": criteria, "sections": use_sections, "drafts": drafts}
        if save_name.strip():
            with open(save_name.strip(),"w",encoding="utf-8") as f: json.dump(blob,f,ensure_ascii=False,indent=2)
            st.success(f"Saved to {save_name.strip()}")
        else:
            st.warning("Enter filename like my-grant.json")

    load_file = st.file_uploader("Load project (.json)", type=["json"])
    if load_file is not None:
        try:
            data = json.load(load_file)
            st.session_state.project = Project(**data.get("project",{}))
            st.session_state.application = data.get("drafts",{})
            st.session_state.loaded_criteria = data.get("criteria","")
            st.session_state.loaded_sections = data.get("sections",SECTIONS)
            st.success("Project loaded.")
        except Exception as e:
            st.error(f"Failed to load JSON: {e}")

# ===== FORM =====
if "project" not in st.session_state: st.session_state.project = Project()
p: Project = st.session_state.project

st.markdown("### 🧾 Project Details")
c1,c2 = st.columns(2)
with c1:
    p.title = st.text_input("Project Title", value=p.title)
    p.summary = st.text_area("Short Summary (2–3 sentences)", height=90, value=p.summary)
    p.need = st.text_area("Problem / Need", height=120, value=p.need)
    p.audience = st.text_area("Target Audience", height=100, value=p.audience)
    p.objectives = st.text_area("Objectives (bullets OK)", height=120, value=p.objectives)
with c2:
    p.activities = st.text_area("Activities & Delivery Plan", height=120, value=p.activities)
    p.outcomes = st.text_area("Expected Outcomes (KPIs/metrics)", height=100, value=p.outcomes)
    p.evaluation = st.text_area("Evaluation (how you'll measure)", height=100, value=p.evaluation)
    p.budget = st.text_area("Budget (summary + justification)", height=120, value=p.budget)
    p.risks = st.text_area("Risks & Mitigation", height=100, value=p.risks)
p.timeline = st.text_input("High-level Timeline", value=p.timeline)
p.partners = st.text_area("Partners & Governance", height=90, value=p.partners)

if "loaded_criteria" in st.session_state and st.session_state.loaded_criteria and not criteria.strip():
    criteria = st.session_state.loaded_criteria
if "loaded_sections" in st.session_state and st.session_state.loaded_sections and use_sections == SECTIONS:
    use_sections = st.session_state.loaded_sections

st.markdown("---")

# ===== GENERATE =====
if "application" not in st.session_state: st.session_state.application = {}
logo_bytes = load_logo_bytes()

if st.button("🚀 Generate Draft + Score Sections", type="primary", use_container_width=True):
    if not p.title or not p.summary:
        st.warning("Please add at least Project Title and Short Summary.")
    else:
        out: Dict[str,str] = {}
        prog = st.progress(0.0, text="Generating…")
        for i, sec in enumerate(use_sections, start=1):
            with st.spinner(f"Generating: {sec}"):
                out[sec] = gpt_text(draft_prompt(sec, p, criteria), model=model)
            prog.progress(i/len(use_sections), text=f"Generated {i}/{len(use_sections)}")
        st.session_state.application = out
        st.success("Draft generated. Scroll down to review & export.")

# ===== REVIEW / SCORE / EXPORT =====
drafts: Dict[str,str] = st.session_state.get("application", {})
if drafts:
    st.markdown("### ✍️ Review, Score & Export")
    edited: Dict[str,str] = {}
    tabs = st.tabs([f"{i+1}. {k}" for i,k in enumerate(drafts)])
    for (k,v), tb in zip(drafts.items(), tabs):
        with tb:
            new = st.text_area(f"Edit: {k}", value=v, height=240, key=f"edit_{k}")
            if st.button(f"Score '{k}'", key=f"score_{k}"):
                res = gpt_json(score_prompt(new, criteria), model=model)
                st.session_state[f"sc_{k}"] = res
            res = st.session_state.get(f"sc_{k}")
            if res:
                st.write(f"**Score:** {res.get('score',0)}/100")
                c = st.columns(3)
                c[0].write("**Strengths**"); c[0].write("\n".join(res.get("strengths",[])) or "—")
                c[1].write("**Gaps**"); c[1].write("\n".join(res.get("gaps",[])) or "—")
                c[2].write("**Suggestions**"); c[2].write("\n".join(res.get("suggestions",[])) or "—")
            edited[k] = new

    st.markdown("---")
    cE1, cE2 = st.columns(2)
    with cE1:
        if st.button("⬇️ Download DOCX", use_container_width=True):
            b = build_docx(edited, p, logo_bytes)
            st.download_button("Download Grant Draft (DOCX)", data=b,
                file_name=f"{COUNCIL_SHORT}_grant_draft_{datetime.datetime.now().strftime('%Y%m%d')}.docx",
                mime="application/vnd.openxmlformats-officedocument.wordprocessingml.document")
    with cE2:
        md = "# "+COUNCIL_NAME+" – Grant Application Draft\n\n" + "\n".join([f"## {k}\n{v}\n" for k,v in edited.items()])
        st.download_button("⬇️ Download Markdown", data=md.encode("utf-8"),
            file_name=f"{COUNCIL_SHORT}_grant_draft_{datetime.datetime.now().strftime('%Y%m%d')}.md", mime="text/markdown")
else:
    st.info("Fill details, paste criteria, then click **Generate**.")
